# ============================================
# KarbonAT - RAPORLAR (raporlama alt sistemi)
#
# Kullanıcının referans-sablonlari/ olarak attığı TGA formatlarının
# HER BİRİ ayrı bir şablondur. AI şablonları verilere göre AYRI AYRI
# doldurur; deterministik olanlar (Tablo 10-13) kodla üretilir.
# Üretilen raporlar data_store üzerinden dönem bazında kaydedilir
# (profil -> raporlar -> aylar).
# ============================================
from __future__ import annotations

import os
import re
from io import BytesIO

import pandas as pd

# Deterministik tablo üreticileri
from tga_tables import (
    tablo10_elektrik,
    tablo11_enerji,
    tablo12_su,
    tablo13_atik,
    kimyasal_envanter,
)

RAPOR_SABLONLARI = [
    {
        "id": "tablo1",
        "emoji": "⚠️",
        "baslik": "Risk Analizi (Tablo 1)",
        "tip": "ai",
        "cikti": ["PDF"],
        "kaynak": ["Tablo 1"],
        "aciklama": "Risk matrisi formatında tesis risklerini veriyle ilişkilendirir; olasılık/etki skorları önerir.",
    },
    {
        "id": "tablo2",
        "emoji": "🎯",
        "baslik": "Hedefler (Tablo 2)",
        "tip": "ai",
        "cikti": ["PDF"],
        "kaynak": ["Tablo 2"],
        "aciklama": "TGA Tablo 2 formatında ölçülebilir sürdürülebilirlik hedefleri; mevcut veriye dayalı baz + hedef.",
    },
    {
        "id": "tablo3",
        "emoji": "⚖️",
        "baslik": "Yasal Uygunluk & Yükümlülükler (Tablo 3)",
        "tip": "ai",
        "cikti": ["PDF"],
        "kaynak": ["Tablo 3"],
        "aciklama": "Çevre ve turizm mevzuatına uygunluk listesi; ilgili yükümlülükleri ve uyum durumu.",
    },
    {
        "id": "anlati",
        "emoji": "📄",
        "baslik": "Sürdürülebilirlik Raporu (Tablo 4)",
        "tip": "ai",
        "cikti": ["PDF"],
        "kaynak": ["Raporlamasi", "Politika"],
        "aciklama": "TGA Tablo 4 formatında, tesis verileriyle doldurulmuş anlatı raporu; yönetim mesajı, performans, hedefler ve uyum maddeleri.",
    },
    {
        "id": "tablo5",
        "emoji": "📊",
        "baslik": "Misafir Anketi Soruları (Tablo 5)",
        "tip": "ai",
        "cikti": ["PDF"],
        "kaynak": ["Tablo 5"],
        "aciklama": "TGA Tablo 5 sürdürülebilirlik anketi soru seti; tesise uyarlanmış memnuniyet + farkındalık soruları.",
    },
    {
        "id": "tablo7",
        "emoji": "🤝",
        "baslik": "Tedarikçi Değerlendirme (Tablo 7)",
        "tip": "ai",
        "cikti": ["PDF"],
        "kaynak": ["Tedarikci"],
        "aciklama": "Tedarikçi değerlendirme formu; sürdürülebilirlik kriterlerine göre skorlama maddeleri.",
    },
    {
        "id": "tablo8",
        "emoji": "✅",
        "baslik": "Onaylı Tedarikçi Listesi (Tablo 8)",
        "tip": "ai",
        "cikti": ["PDF"],
        "kaynak": ["Tablo 8"],
        "aciklama": "Onaylı tedarikçi listesi; kriterler ve onay tarihi sütunlarıyla doldurulmuş taslak.",
    },
    {
        "id": "tablo9",
        "emoji": "♻️",
        "baslik": "Tek Kullanımlık Plastik Sarfiyatı (Tablo 9)",
        "tip": "ai",
        "cikti": ["PDF"],
        "kaynak": ["Tablo 9"],
        "aciklama": "Plastik/ambalajlı ürün sarfiyat takip formatı; azaltma önerileriyle.",
    },
    {
        "id": "tablo10",
        "emoji": "🔌",
        "baslik": "Elektrik Tüketim Takibi (Tablo 10)",
        "tip": "deterministik",
        "cikti": ["Excel"],
        "kaynak": [],
        "aciklama": "Aylık elektrik tüketimi ve emisyon (kWh → kg CO₂e).",
    },
    {
        "id": "tablo11",
        "emoji": "🔥",
        "baslik": "Enerji / Yakıt Takibi (Tablo 11)",
        "tip": "deterministik",
        "cikti": ["Excel"],
        "kaynak": [],
        "aciklama": "Doğal gaz ve yakıt tüketimi; yenilenebilir oranı.",
    },
    {
        "id": "tablo12",
        "emoji": "🚿",
        "baslik": "Su Sarfiyatı Takibi (Tablo 12)",
        "tip": "deterministik",
        "cikti": ["Excel"],
        "kaynak": [],
        "aciklama": "Aylık su kullanımı; kişi ve oda-gün başına normalizasyon.",
    },
    {
        "id": "tablo13",
        "emoji": "🗑️",
        "baslik": "Atık Takibi (Tablo 13)",
        "tip": "deterministik",
        "cikti": ["Excel"],
        "kaynak": [],
        "aciklama": "Atık türü bazında aylık miktar; bertaraf ve geri kazanım.",
    },
    {
        "id": "politika_set",
        "emoji": "📋",
        "baslik": "Sürdürülebilirlik Politika Seti",
        "tip": "ai",
        "cikti": ["PDF"],
        "kaynak": ["POLITIKASI", "Politika"],
        "aciklama": "Çevre/atık, enerji, kadın & cinsiyet, çocuk hakları, satın alma ve ana politika belgesi; her biri kendi maddeleriyle tesise uyarlanır.",
    },
]


def sablon_bul(sablon_id: str) -> dict | None:
    for s in RAPOR_SABLONLARI:
        if s["id"] == sablon_id:
            return s
    return None


# ---------------- Deterministik (Tablo 10-13) ----------------
def _df_md(df) -> str:
    """DataFrame'i basit markdown tablosuna çevirir (tabulate bağımlılığı yok)."""
    import pandas as pd

    if df is None or len(df) == 0:
        return "(veri yok)"
    sutunlar = [str(c) for c in df.columns]
    satirlar = [[("" if pd.isna(v) else str(v)) for v in row] for row in df.itertuples(index=False)]
    genislik = {i: max(len(s), *(len(r[i]) for r in satirlar)) for i, s in enumerate(sutunlar)}
    ayrac = "| " + " | ".join("-" * genislik[i] for i in range(len(sutunlar))) + " |"
    baslik = "| " + " | ".join(sutunlar[i].ljust(genislik[i]) for i in range(len(sutunlar))) + " |"
    govde = "\n".join(
        "| " + " | ".join(r[i].ljust(genislik[i]) for i in range(len(sutunlar))) + " |"
        for r in satirlar
    )
    return f"{baslik}\n{ayrac}\n{govde}"


def _deterministik_xlsx(sablon_id: str, sonuc: dict) -> bytes:
    tuketim = sonuc["tuketim"]
    if sablon_id == "tablo10":
        df = tablo10_elektrik(tuketim)
    elif sablon_id == "tablo11":
        df = tablo11_enerji(tuketim)
    elif sablon_id == "tablo12":
        tesis = sonuc["tesis"]
        df = tablo12_su(tuketim, tesis.get("dolu_oda_gun", 0), tesis.get("musteri", 0))
    elif sablon_id == "tablo13":
        df = tablo13_atik(tuketim)
    else:
        raise ValueError(f"Bilinmeyen deterministik şablon: {sablon_id}")

    buffer = BytesIO()
    with __import__("pandas").ExcelWriter(buffer, engine="openpyxl") as writer:
        df.to_excel(writer, sheet_name="Tablo", index=False)
        ws = writer.sheets["Tablo"]
        # Denetime hazır stil – aynı helper
        try:
            _xlsx_stil_uygula(ws, df)
        except Exception:
            pass
    return buffer.getvalue()


def _deterministik_md(sablon_id: str, sonuc: dict) -> str:
    tuketim = sonuc["tuketim"]
    if sablon_id == "tablo10":
        df = tablo10_elektrik(tuketim)
    elif sablon_id == "tablo11":
        df = tablo11_enerji(tuketim)
    elif sablon_id == "tablo12":
        tesis = sonuc["tesis"]
        df = tablo12_su(tuketim, tesis.get("dolu_oda_gun", 0), tesis.get("musteri", 0))
    elif sablon_id == "tablo13":
        df = tablo13_atik(tuketim)
    else:
        return ""
    return _df_md(df)


# ---------------- AI şablonları ----------------
def _ai_uretim(sablon: dict, tesis: dict, sonuc: dict | None, prefs: dict) -> str:
    import ai_engine

    return ai_engine.oner_rapor(sablon, tesis, sonuc, prefs)


# ---------------- Ana API ----------------
def rapor_uretim(sablon_id: str, tesis: dict, sonuc: dict | None, prefs: dict | None = None) -> dict:
    """Rapor üretir. Dönen: {"sablon_id", "tip", "metin", "xlsx": bytes|None}"""
    sablon = sablon_bul(sablon_id)
    if not sablon:
        raise ValueError(f"Bilinmeyen rapor şablonu: {sablon_id}")
    prefs = prefs or {}
    if sablon["tip"] == "deterministik":
        metin = _deterministik_md(sablon_id, sonuc or {})
        xlsx = _deterministik_xlsx(sablon_id, sonuc or {})
        return {"sablon_id": sablon_id, "tip": "deterministik", "metin": metin, "xlsx": xlsx}
    metin = _ai_uretim(sablon, tesis, sonuc, prefs)
    return {"sablon_id": sablon_id, "tip": "ai", "metin": metin, "xlsx": None}


# ---------------- Markdown tablo ayrıştırıcı (UI + DOCX + XLSX ortak) ----------------
def markdown_bloklar(metin: str) -> list[tuple]:
    """Markdown'ı bloklara böler: ("md", metin) veya ("tablo", DataFrame)."""
    satirlar = metin.splitlines()
    bloklar = []
    i = 0
    n = len(satirlar)
    while i < n:
        s = satirlar[i].strip()
        if s.startswith("|"):
            tab = []
            while i < n and satirlar[i].strip().startswith("|"):
                tab.append(satirlar[i].strip())
                i += 1
            rows = [[c.strip() for c in r.strip("|").split("|")] for r in tab]
            rows = [r for r in rows if not all(set(c) <= {"-", ":", " "} for c in r)]
            rows = [r for r in rows if any(r)]
            if len(rows) >= 2:
                ncols = max(len(r) for r in rows)
                rows = [r[:ncols] + [""] * (ncols - len(r)) for r in rows]
                cols = []
                gorulen = set()
                for j, c in enumerate(rows[0]):
                    ad = c.strip()
                    if not ad:
                        ad = f"Sütun {j + 1}"
                    if ad in gorulen:
                        ad = f"{ad} ({j + 1})"
                    gorulen.add(ad)
                    cols.append(ad)
                df = pd.DataFrame(rows[1:], columns=cols)
                bloklar.append(("tablo", df))
            elif rows:
                bloklar.append(("md", rows[0][0]))
        elif not s:
            i += 1
        else:
            parca = []
            while i < n:
                k = satirlar[i].strip()
                if not k or k.startswith("|"):
                    break
                parca.append(satirlar[i])
                i += 1
            bloklar.append(("md", "\n".join(parca)))
    return bloklar


# ---------------- DOCX (gerçek Word tabloları) ----------------
def _docx_ekle_blok(doc, blok):
    tur, icerik = blok
    if tur == "tablo":
        df = icerik
        if doc.tables:
            doc.add_paragraph()
        ncols = max(1, len(df.columns))
        tablo = doc.add_table(rows=1, cols=ncols)
        tablo.style = "Light Grid Accent 1"
        tablo.autofit = True
        # Dar kenar boşlukları – geniş tablo sığsın
        try:
            from docx.shared import Inches
            for sec in doc.sections:
                sec.left_margin = Inches(0.5)
                sec.right_margin = Inches(0.5)
                sec.top_margin = Inches(0.5)
                sec.bottom_margin = Inches(0.5)
        except Exception:
            pass
        # Başlık satırı stil
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        header_cells = tablo.rows[0].cells
        for j, c in enumerate(df.columns):
            cell = header_cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(c))
            run.bold = True
            run.font.size = Pt(7.5 if ncols >= 8 else 8.5)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            # Arkaplan
            shading = cell._element.get_or_add_tcPr()
            shd = shading.makeelement(qn('w:shd'), {
                qn('w:fill'): '1d6b45', qn('w:val'): 'clear'
            })
            shading.append(shd)
            # Hücre kenar boşluğu
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            for pp in cell.paragraphs:
                pp.paragraph_format.space_after = Pt(2)
        # Veri satırları – zebra
        for idx, (_, row) in enumerate(df.iterrows()):
            hucreler = tablo.add_row().cells
            for j, v in enumerate(row):
                if j < len(hucreler):
                    cell = hucreler[j]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    txt = "" if pd.isna(v) else str(v)
                    run = p.add_run(txt)
                    run.font.size = Pt(7 if ncols >= 8 else 8)
                    run.font.color.rgb = RGBColor(0x1A, 0x1F, 0x1C)
                    # zebra arka plan
                    if idx % 2 == 1:
                        shading = cell._element.get_or_add_tcPr()
                        shd = shading.makeelement(qn('w:shd'), {
                            qn('w:fill'): 'eef4ee', qn('w:val'): 'clear'
                        })
                        shading.append(shd)
                    # Wrap
                    p.paragraph_format.space_after = Pt(1)
        # Kolon genişlikleri – ağırlıklı dağıtım (docx autofit ile orantılı)
        # docx autofit True olduğundan manuel genişlik ayarı opsiyonel; en uzun metne göre orantı için width hint verelim
        try:
            from docx.shared import Inches
            total_in = 7.5  # kullanılabilir genişlik
            weights = []
            for col in df.columns:
                hlen = len(str(col))
                max_cell = max((len(str(v)) for v in df[col].astype(str)), default=0) if len(df) else 0
                w = max(hlen, max_cell, 4)
                w = min(w, 28)
                weights.append(w)
            tot = sum(weights) or 1
            for i, w in enumerate(weights):
                # width hint
                target = total_in * w / tot
                # En dar sütun 0.6 in altına düşmesin
                target = max(target, 0.55)
                for row in tablo.rows:
                    row.cells[i].width = Inches(target)
        except Exception:
            pass
        return
    for satir in icerik.splitlines():
        s = satir.strip()
        if not s:
            continue
        if s in ("---", "***", "___") or set(s) <= {"-", "*", "_"}:
            continue
        if s.startswith("### "):
            doc.add_heading(s[4:].replace("**", ""), level=2)
        elif s.startswith("## "):
            doc.add_heading(s[3:].replace("**", ""), level=1)
        elif s.startswith("# "):
            doc.add_heading(s[2:].replace("**", ""), level=0)
        elif s.startswith("- ") or s.startswith("* ") or s.startswith("• "):
            txt = s.lstrip("-*• ").replace("**", "").strip()
            if not txt or set(txt) <= {"-", "–", "—"}:
                continue
            doc.add_paragraph(txt, style="List Bullet")
        else:
            doc.add_paragraph(s.replace("**", ""))


def rapor_docx(metin: str) -> bytes:
    """Markdown'ı gerçek Word tablolarıyla .docx yapar – geniş tablolar dar kenar boşluğu + wrap ile sığar."""
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    # Dar kenar boşlukları
    for sec in doc.sections:
        sec.left_margin = Inches(0.5)
        sec.right_margin = Inches(0.5)
        sec.top_margin = Inches(0.6)
        sec.bottom_margin = Inches(0.5)
    # Kurumsal kapak başlığı
    try:
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("KarbonAT P2  ·  GSTC / TGA Uyumlu  ·  Denetime Hazır")
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x1D, 0x6B, 0x45)
        p.paragraph_format.space_after = Pt(6)
        # ince ayırıcı
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run("—" * 48)
        r2.font.size = Pt(6)
        r2.font.color.rgb = RGBColor(0xCF, 0xD8, 0xCF)
        p2.paragraph_format.space_after = Pt(8)
        # Footer için section footer (opsiyonel)
        try:
            footer = sec.footer.paragraphs[0]
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fr = footer.add_run("KarbonAT P2  ·  Gerçek veriden üretildi  ·  TGA/GSTC uyumlu")
            fr.font.size = Pt(7)
            fr.font.color.rgb = RGBColor(0x6B, 0x7A, 0x70)
        except Exception:
            pass
    except Exception:
        pass
    for blok in markdown_bloklar(metin):
        _docx_ekle_blok(doc, blok)
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ---------------- XLSX (tablo başına ayrı sayfa) ----------------
def _xlsx_stil_uygula(ws, df, is_metin=False):
    """openpyxl Worksheet'e denetime hazır stil uygular: başlık, zebra, wrap, filter, freeze."""
    from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    HEADER_FILL = PatternFill(start_color="1D6B45", end_color="1D6B45", fill_type="solid")
    HEADER_FONT = Font(name="Calibri", size=9, bold=True, color="FFFFFF")
    HEADER_ALIGN = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ZEBRA_FILL = PatternFill(start_color="EEF4EE", end_color="EEF4EE", fill_type="solid")
    CELL_ALIGN = Alignment(horizontal="left", vertical="center", wrap_text=True)
    THIN = Side(style="thin", color="CFD8CF")
    BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)

    ncols = len(df.columns) if not is_metin else 1
    nrows = len(df) + 1 if not is_metin else len(df) + 1
    # Başlık satırı
    for col_idx in range(1, ncols + 1):
        cell = ws.cell(row=1, column=col_idx)
        cell.fill = HEADER_FILL
        cell.font = HEADER_FONT
        cell.alignment = HEADER_ALIGN
        cell.border = BORDER
    # Zebra
    for r in range(2, nrows + 1):
        is_zebra = (r % 2 == 0)  # 2. satır beyaz, 3. satır zebra vb – değiştirilebilir
        for c in range(1, ncols + 1):
            cell = ws.cell(row=r, column=c)
            cell.alignment = CELL_ALIGN
            cell.border = BORDER
            cell.font = Font(name="Calibri", size=9 if ncols < 8 else 8)
            if is_zebra and r >= 3 and (r % 2 == 1):
                cell.fill = ZEBRA_FILL
            # Sayısal hücreleri sağa hizala
            try:
                if isinstance(cell.value, (int, float)):
                    cell.alignment = Alignment(horizontal="right", vertical="center")
            except Exception:
                pass
    # Kolon genişlikleri – içerik tabanlı ama sınırlı
    for col_idx, col_name in enumerate(df.columns, 1):
        # En uzun metin
        max_len = len(str(col_name))
        for val in df[col_name].astype(str):
            if len(val) > max_len:
                max_len = len(val)
        # Genişlik = char * 1.1 + padding, 12-38 arası
        width = min(38, max(12, max_len * 1.05 + 2))
        # Çok sütunlu tablolarda genişlikleri küçült
        if ncols >= 10:
            width = min(width, 18)
        elif ncols >= 7:
            width = min(width, 24)
        ws.column_dimensions[get_column_letter(col_idx)].width = width
    # Yükseklik – başlık daha yüksek (wrap için)
    ws.row_dimensions[1].height = 28 if ncols >= 8 else 20
    for r in range(2, nrows + 1):
        ws.row_dimensions[r].height = 16
    # Freeze + Filter
    ws.freeze_panes = "A2"
    try:
        ws.auto_filter.ref = ws.dimensions
    except Exception:
        pass
    # Yazdırma ayarları
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.page_setup.orientation = "landscape" if ncols >= 7 else "portrait"
    ws.page_setup.paperSize = ws.PAPERSIZE_A4
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.page_margins.left = 0.4
    ws.page_margins.right = 0.4
    ws.page_margins.top = 0.5
    ws.page_margins.bottom = 0.5
    ws.print_title_rows = "1:1"


def rapor_xlsx(metin: str) -> bytes:
    """Markdown tablolarını ayrı sayfalara; kalan metni 'Metin' sayfasına yazar – kurumsal Kapak + styled & printable."""
    from datetime import datetime
    buffer = BytesIO()
    bloklar = markdown_bloklar(metin)
    tablolar = [b for b in bloklar if b[0] == "tablo"]
    metinler = [b[1] for b in bloklar if b[0] == "md"]
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        # --- Kapak sayfası (kurumsal kimlik) ---
        from openpyxl.styles import Font, Alignment, PatternFill
        kapak = pd.DataFrame([
            ["KarbonAT P2 – Rapor Eki (AI)"],
            [f"Oluşturulma: {datetime.now().strftime('%d.%m.%Y %H:%M')}"],
            [""],
            ["Bu dosya KarbonAT tarafından tesisin gerçek verilerinden üretilmiştir."],
            ["Tablolar denetime hazır, filtreli ve yazdırma ayarlıdır."],
            ["Her tablo ayrı sayfadadır; 'Metin' sayfasında anlatı bölümleri yer alır."],
            ["GSTC / TGA uyumlu – KarbonAT P2"],
        ])
        kapak.to_excel(writer, sheet_name="Kapak", index=False, header=False)
        ws_k = writer.sheets["Kapak"]
        ws_k["A1"].font = Font(name="Calibri", size=13, bold=True, color="1D6B45")
        ws_k["A1"].alignment = Alignment(horizontal="left", vertical="center")
        ws_k["A2"].font = Font(name="Calibri", size=9, color="6B7A70")
        for r in range(4, 8):
            ws_k.cell(row=r, column=1).font = Font(name="Calibri", size=9, color="17201c")
            ws_k.cell(row=r, column=1).alignment = Alignment(wrap_text=True, vertical="center")
        ws_k.column_dimensions["A"].width = 78
        ws_k.row_dimensions[1].height = 18
        ws_k.sheet_properties.pageSetUpPr.fitToPage = True
        ws_k.page_setup.orientation = "portrait"
        ws_k.page_setup.paperSize = ws_k.PAPERSIZE_A4
        ws_k.page_setup.fitToWidth = 1
        ws_k.page_margins.left = 0.6
        ws_k.page_margins.right = 0.6

        for i, (_, df) in enumerate(tablolar):
            sheet = f"Tablo_{i + 1}"[:31]
            df.to_excel(writer, sheet_name=sheet, index=False)
            ws = writer.sheets[sheet]
            _xlsx_stil_uygula(ws, df)
        if metinler:
            metin_df = pd.DataFrame({"İçerik": [m for p in metinler for m in p.splitlines() if m.strip()]})
            metin_df.to_excel(writer, sheet_name="Metin", index=False)
            ws = writer.sheets["Metin"]
            _xlsx_stil_uygula(ws, metin_df, is_metin=True)
            ws.column_dimensions["A"].width = 90
            ws.column_dimensions["A"].width = 85
        if not tablolar and not metinler:
            pd.DataFrame({"Bilgi": ["İçerik yok – AI raporu boş"]}).to_excel(writer, sheet_name="Bos", index=False)
    return buffer.getvalue()


# ---------------- PDF (Türkçe font destekli) ----------------
_PDF_FONTLAR = None


def _pdf_font_ayar():
    """Türkçe karakterleri gösteren bir TTF font bulup reportlab'e kaydeder."""
    global _PDF_FONTLAR
    if _PDF_FONTLAR:
        return _PDF_FONTLAR
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # Önce proje içindeki Roboto (repo ile birlikte gelir, Linux'ta da Türkçe destekler)
    proj_root = os.path.dirname(os.path.abspath(__file__))
    adaylar = [
        ("Roboto", os.path.join(proj_root, "Roboto-Regular.ttf"), "Roboto-Bold", os.path.join(proj_root, "Roboto-Regular.ttf")),
        ("RobotoData", os.path.join(proj_root, "data", "Roboto-Regular.ttf"), "RobotoData-Bold", os.path.join(proj_root, "data", "Roboto-Regular.ttf")),
        ("ArialTR", r"C:\Windows\Fonts\arial.ttf", "ArialTR-Bold", r"C:\Windows\Fonts\arialbd.ttf"),
        ("DejaVuSans", r"C:\Windows\Fonts\dejavusans.ttf", "DejaVuSans-Bold", r"C:\Windows\Fonts\dejavusans-bold.ttf"),
        ("TimesNewRomanTR", r"C:\Windows\Fonts\times.ttf", "TimesNewRomanTR-Bold", r"C:\Windows\Fonts\timesbd.ttf"),
    ]
    for ad, yol, bad, byol in adaylar:
        if os.path.exists(yol):
            pdfmetrics.registerFont(TTFont(ad, yol))
            if os.path.exists(byol):
                try:
                    pdfmetrics.registerFont(TTFont(bad, byol))
                except Exception:  # noqa: BLE001
                    bad = None
            _PDF_FONTLAR = (ad, bad)
            return _PDF_FONTLAR
    _PDF_FONTLAR = ("Helvetica", "Helvetica-Bold")
    return _PDF_FONTLAR


def _pdf_table(df, avail_w, font_ad, font_bold):
    """DataFrame'i sığan, wrap'li bir ReportLab Table'a çevirir."""
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import Paragraph, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    ncols = len(df.columns)
    # Font boyutu – sütun sayısına göre küçült
    if ncols <= 5:
        fs, hfs = 7.5, 8.0
    elif ncols <= 7:
        fs, hfs = 6.8, 7.2
    elif ncols <= 10:
        fs, hfs = 6.0, 6.5
    else:
        fs, hfs = 5.4, 5.8

    # Sütun ağırlıkları – uzun metni sınırlayarak orantılı dağıtım
    weights = []
    for col in df.columns:
        hlen = len(str(col))
        max_cell = 0
        if len(df):
            try:
                max_cell = max(len(str(v)) for v in df[col].astype(str) if str(v) != "nan")
            except Exception:
                max_cell = hlen
        w = max(hlen, max_cell, 4)
        w = min(w, 30)  # çok uzun metin ağırlığı sınırlı – wrap ile sığar
        # Dar sütunlar (No, Etki, Olasılık) için minimum ağırlık
        if hlen <= 3:
            w = max(w, 4)
        weights.append(w)
    tot = sum(weights) or 1
    col_widths = [avail_w * w / tot for w in weights]
    # Minimum genişlik koruması: her sütun en az 28pt olsun
    col_widths = [max(cw, 28) for cw in col_widths]
    # Toplam avail_w'yi aşarsa orantılı küçült
    s = sum(col_widths)
    if s > avail_w:
        scale = avail_w / s
        col_widths = [cw * scale for cw in col_widths]

    # Hücre stilleri – wrap'li Paragraph
    cell_style = ParagraphStyle("cell", fontName=font_ad, fontSize=fs, leading=fs * 1.25 + 0.5,
                                textColor=colors.HexColor("#1a1f1c"), alignment=TA_LEFT,
                                spaceBefore=1, spaceAfter=1, wordWrap="CJK")
    header_style = ParagraphStyle("header", fontName=font_bold, fontSize=hfs, leading=hfs * 1.2 + 0.5,
                                  textColor=colors.white, alignment=TA_CENTER,
                                  spaceBefore=1, spaceAfter=1, wordWrap="CJK")

    def _escape(s):
        return str(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    header_row = [Paragraph(_escape(str(c)), header_style) for c in df.columns]
    data = [header_row]
    for _, row in df.iterrows():
        data.append([Paragraph(_escape("" if pd.isna(v) else str(v)), cell_style) for v in row])

    t = Table(data, colWidths=col_widths, repeatRows=1, hAlign="LEFT")
    # Zarif stil
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1d6b45")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#cfd8cf")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#eef4ee")]),
        ("LEFTPADDING", (0, 0), (-1, -1), 3),
        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ("LINEBELOW", (0, 0), (-1, 0), 0.8, colors.HexColor("#143a28")),
    ]))
    return t


def rapor_pdf(metin: str) -> bytes:
    """Rapor metnini Türkçe karakter destekli PDF'e çevirir; tablolar otomatik sığar (wrap + landscape + font adapt)."""
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    from reportlab.lib import colors

    font_ad, font_bold = _pdf_font_ayar()
    bold = font_bold or font_ad
    buffer = BytesIO()

    # Sayfa yönünü tablolara göre seç: >=8 sütun ise yatay daha ferah
    bloklar = markdown_bloklar(metin)
    max_cols = 0
    for tur, ic in bloklar:
        if tur == "tablo":
            max_cols = max(max_cols, len(ic.columns))
    pagesize = landscape(A4) if max_cols >= 8 else A4

    stiller = getSampleStyleSheet()
    baslik = ParagraphStyle("Bas", parent=stiller["Title"], fontName=bold, fontSize=14, spaceAfter=8,
                            textColor=colors.HexColor("#1d6b45"), leading=16)
    h2 = ParagraphStyle("H2", parent=stiller["Heading2"], fontName=bold, fontSize=11.5, spaceBefore=6, spaceAfter=3,
                        textColor=colors.HexColor("#143a28"), leading=14)
    h3 = ParagraphStyle("H3", parent=stiller["Heading3"], fontName=bold, fontSize=10, spaceBefore=5, spaceAfter=2,
                        textColor=colors.HexColor("#2e8b57"), leading=12)
    govde = ParagraphStyle("Govde", parent=stiller["BodyText"], fontName=font_ad, fontSize=9, leading=13,
                           textColor=colors.HexColor("#1a1f1c"), spaceBefore=2, spaceAfter=2)
    bullet = ParagraphStyle("Bullet", parent=govde, leftIndent=14, firstLineIndent=0, spaceBefore=2, spaceAfter=2,
                            bulletIndent=7, alignment=0)

    # Kurumsal header + footer (her sayfa)
    def _header_footer(canvas, doc):
        canvas.saveState()
        w, h = pagesize
        # Header band – kurumsal kimlik
        canvas.setFillColor(colors.HexColor("#1d6b45"))
        canvas.rect(0, h - 9*mm, w, 9*mm, fill=1, stroke=0)
        canvas.setFillColor(colors.white)
        canvas.setFont(bold, 9)
        canvas.drawString(12*mm, h - 6*mm, "KarbonAT P2")
        canvas.setFont(font_ad, 6.5)
        # Sagda ince slogan
        canvas.setFillColor(colors.HexColor("#c8e6c9"))
        canvas.drawRightString(w - 12*mm, h - 6*mm, "GSTC / TGA Uyumlu  ·  Denetime Hazir")
        # Footer
        canvas.setFillColor(colors.HexColor("#6b7a70"))
        canvas.setFont(font_ad, 6.5)
        canvas.drawString(12*mm, 10*mm, "KarbonAT P2  ·  GSTC / TGA uyumlu  ·  Gercek veriden uretildi")
        canvas.drawRightString(w - 12*mm, 10*mm, f"Sayfa {doc.page}")
        canvas.setStrokeColor(colors.HexColor("#e3e9e3"))
        canvas.setLineWidth(0.35)
        canvas.line(12*mm, 12*mm, w - 12*mm, 12*mm)
        canvas.restoreState()

    doc = SimpleDocTemplate(buffer, pagesize=pagesize, topMargin=16*mm, bottomMargin=14*mm,
                            leftMargin=12*mm, rightMargin=12*mm,
                            title="KarbonAT Rapor", author="KarbonAT P2")
    avail_w = doc.width
    akis = []
    for tur, icerik in bloklar:
        if tur == "tablo":
            df = icerik
            t = _pdf_table(df, avail_w, font_ad, bold)
            akis.append(t)
            akis.append(Spacer(1, 7))
            continue
        for satir in icerik.splitlines():
            s = satir.strip()
            if not s:
                akis.append(Spacer(1, 3))
                continue
            # Sadece '---' gibi ayraçları bullet sanma
            if s in ("---", "***", "___") or set(s) <= {"-", "*", "_"}:
                akis.append(Spacer(1, 4))
                continue
            if s.startswith("### "):
                akis.append(Paragraph(s[4:].replace("**", ""), h3))
            elif s.startswith("## "):
                akis.append(Paragraph(s[3:].replace("**", ""), h2))
            elif s.startswith("# "):
                akis.append(Paragraph(s[2:].replace("**", ""), baslik))
            elif s.startswith("- ") or s.startswith("* ") or s.startswith("• "):
                txt = s.lstrip("-*• ").replace("**", "").strip()
                if not txt or set(txt) <= {"-", "–", "—"}:
                    continue
                # Kurumsal bullet – girintili
                akis.append(Paragraph(txt, bullet, bulletText="•"))
            else:
                akis.append(Paragraph(s.replace("**", ""), govde))
    doc.build(akis, onFirstPage=_header_footer, onLaterPages=_header_footer)
    return buffer.getvalue()
