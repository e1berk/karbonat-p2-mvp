# ============================================
# KarbonAT - RAG Bilgi Bankası Derleyicisi (A4)
#
# referans-sablonlari/ içindeki TGA şablonlarını (politika docx'leri,
# eğitim formları, Tablo 1-13) metne çevirir, parçalara (chunk) böler,
# Gemini embedding (gemini-embedding-001) ile vektörleştirir ve
# data/kb.json içine kaydeder. İçerik üretimi (A5) bu dosyayı RAG
# kaynağı olarak kullanır.
#
# Kullanım:  python kb_build.py
#
# NOT: .doc (eski Word) ikili formatı şimdilik atlanır; .docx/.xlsx/.xls
# tam desteklenir. Llama karbonat.docx ve GPT.docx bilinçli dışarıda.
# ============================================
from __future__ import annotations

import json
import os
import re
import sys
from pathlib import Path

from dotenv import load_dotenv

load_dotenv()

KB_KLASOR = Path(__file__).parent / "referans-sablonlari"
CAKTI_YOL = Path(__file__).parent / "data" / "kb.json"
EMBED_MODEL = "gemini-embedding-001"
CHUNK_HEDEF = 1000
CHUNK_ORTUSME = 120
BATCH = 50
DISARIDA = {"llama karbonat.docx", "GPT.docx"}


def _client():
    from google import genai

    anahtar = os.environ.get("GEMINI_API_KEY", "")
    if not anahtar:
        raise RuntimeError("GEMINI_API_KEY .env içinde yok.")
    return genai.Client(api_key=anahtar)


# ---------------- Metin çıkarma ----------------
def _docx_metin(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    bolumler = []
    for p in doc.paragraphs:
        if p.text.strip():
            bolumler.append(p.text.strip())
    for tablo in doc.tables:
        for satir in tablo.rows:
            hucreler = [c.text.strip().replace("\n", " ") for c in satir.cells]
            if any(hucreler):
                bolumler.append(" | ".join(hucreler))
    return "\n".join(bolumler)


def _xlsx_metin(path: Path) -> str:
    import openpyxl

    wb = openpyxl.load_workbook(str(path), data_only=True)
    bolumler = [f"[SHEET: {ws.title}]" for ws in wb.worksheets]
    for ws in wb.worksheets:
        for satir in ws.iter_rows(values_only=True):
            degerler = ["" if v is None else str(v).strip().replace("\n", " ") for v in satir]
            if any(degerler):
                bolumler.append(" | ".join(degerler))
    return "\n".join(bolumler)


def _xls_metin(path: Path) -> str:
    import xlrd

    wb = xlrd.open_workbook(str(path))
    bolumler = []
    for ws in wb.sheets():
        bolumler.append(f"[SHEET: {ws.name}]")
        for r in range(ws.nrows):
            degerler = [str(ws.cell_value(r, c)).strip() for c in range(ws.ncols)]
            if any(degerler):
                bolumler.append(" | ".join(degerler))
    return "\n".join(bolumler)


_CIKARICILAR = {
    ".docx": _docx_metin,
    ".xlsx": _xlsx_metin,
    ".xls": _xls_metin,
}


def _dosya_metni(path: Path) -> str:
    uzanti = path.suffix.lower()
    if uzanti not in _CIKARICILAR:
        print(f"  ! {path.name}: desteklenmeyen uzantı ({uzanti}), atlandı")
        return ""
    try:
        return _CIKARICILAR[uzanti](path)
    except Exception as e:
        print(f"  ! {path.name}: okunamadı ({e})")
        return ""


# ---------------- Parçalama ----------------
def _cumlelere_bol(metin: str) -> list[str]:
    parc = re.split(r"(?<=[.?!])\s+|\n+", metin)
    return [p.strip() for p in parc if len(p.strip()) > 8]


def _chunkla(metin: str) -> list[str]:
    cumleler = _cumlelere_bol(metin)
    chunklar: list[str] = []
    mevcut = ""
    for c in cumleler:
        if len(mevcut) + len(c) + 1 <= CHUNK_HEDEF:
            mevcut = f"{mevcut} {c}".strip()
        else:
            if mevcut:
                chunklar.append(mevcut)
            # Çok uzun cümle yine de parçalansın
            while len(c) > CHUNK_HEDEF:
                chunklar.append(c[:CHUNK_HEDEF])
                c = c[CHUNK_HEDEF:]
            mevcut = c
    if mevcut:
        chunklar.append(mevcut)
    return chunklar


# ---------------- Embedding ----------------
def _embed(c: object, metinler: list[str]) -> list[list[float]]:
    sonuc = c.models.embed_content(
        model=EMBED_MODEL,
        contents=metinler,
    )
    return [e.values for e in sonuc.embeddings]


def _build() -> dict:
    c = _client()
    belgeler = []
    for path in sorted(KB_KLASOR.iterdir()):
        if path.name in DISARIDA or not path.is_file():
            continue
        metin = _dosya_metni(path)
        if not metin:
            continue
        chunklar = _chunkla(metin)
        belgeler.append((path.name, chunklar))
        print(f"  + {path.name}: {len(chunklar)} chunk")

    toplam_chunk = sum(len(ch) for _, ch in belgeler)
    print(f"Toplam {toplam_chunk} chunk vektörleniyor...")

    kayitlar = []
    pid = 0
    for kaynak, chunklar in belgeler:
        for i in range(0, len(chunklar), BATCH):
            parti = chunklar[i : i + BATCH]
            embeds = _embed(c, parti)
            for chunk, emb in zip(parti, embeds):
                kayitlar.append(
                    {"id": f"{pid}", "kaynak": kaynak, "metin": chunk, "embedding": emb}
                )
                pid += 1
            print(f"  ~ {pid}/{toplam_chunk}")

    CAKTI_YOL.parent.mkdir(parents=True, exist_ok=True)
    with open(CAKTI_YOL, "w", encoding="utf-8") as f:
        json.dump(
            {"model": EMBED_MODEL, "boyut": len(kayitlar[0]["embedding"]) if kayitlar else 0,
             "kaynak_sayisi": len(belgeler), "chunks": kayitlar},
            f, ensure_ascii=False,
        )
    print(f"OK {CAKTI_YOL} kaydedildi ({len(kayitlar)} chunk).")
    return kayitlar


if __name__ == "__main__":
    _build()
